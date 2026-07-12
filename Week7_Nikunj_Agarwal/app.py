import os
os.environ["ANONYMIZED_TELEMETRY"] = "False"

import streamlit as st
import chromadb
from chromadb.utils import embedding_functions
from dotenv import load_dotenv
from groq import Groq
import pypdf
import docx

# ---------- Setup ----------
@st.cache_resource
def load_collection():
    chroma_client = chromadb.PersistentClient(path="chroma_db")
    ef = embedding_functions.DefaultEmbeddingFunction()
    return chroma_client.get_or_create_collection(
        name="policy_docs",
        embedding_function=ef
    )

collection = load_collection()
load_dotenv()
groq_client = Groq()

def simple_chunk_text(text, chunk_size=300, overlap=50):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

def ask_policy_bot(question, model_name):
    search_results = collection.query(query_texts=[question], n_results=8)
    relevant_chunks = search_results["documents"][0]
    context = "\n\n".join(relevant_chunks)

    prompt = f"""You are a professional HR assistant. Answer the question using ONLY the context below.

Guidelines:
- Give a clear, complete answer in 2-4 sentences.
- Be professional and to the point — no unnecessary elaboration or filler.
- Include relevant specific details (numbers, conditions, processes) directly from the context.
- Do not repeat the question or add generic introductions like "Based on the policy...".
- If the answer is not in the context, simply say "I don't have that information in the policy documents."

Context:
{context}

Question: {question}

Answer:"""

    response = groq_client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=500
    )
    return response.choices[0].message.content

# ---------- Page Setup ----------
st.set_page_config(page_title="PolicyBot", page_icon="📋")
st.title("📋 Saksoft PolicyBot")
st.caption("Ask me anything about company policies — attach new policy docs using the + icon")

model_choice = st.selectbox(
    "Model",
    options=[
        "llama-3.1-8b-instant",
        "llama-3.3-70b-versatile",
        "gemma2-9b-it",
        "openai/gpt-oss-20b"
    ],
    index=0,
    label_visibility="collapsed"
)

# ---------- Chat History ----------
if "messages" not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

# ---------- Chat Input (with attach/+ icon) ----------
user_input = st.chat_input(
    "Ask a question about company policy...",
    accept_file=True,
    file_type=["txt", "pdf", "docx"]
)

if user_input:
    user_question = user_input.text
    uploaded_files = user_input["files"]

    # ---------- File Upload Processing (with live status) ----------
    if uploaded_files:
        for uploaded_file in uploaded_files:
            with st.status(f"Processing '{uploaded_file.name}'...", expanded=True) as status:

                st.write("📁 Saving file...")
                os.makedirs("data/policies", exist_ok=True)
                save_path = os.path.join("data/policies", uploaded_file.name)
                with open(save_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                st.write(f"✅ Saved to `{save_path}`")

                st.write("📖 Reading document content...")
                if uploaded_file.name.endswith(".pdf"):
                    reader = pypdf.PdfReader(save_path)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                    st.write(f"✅ Extracted text from {len(reader.pages)} pages")
                elif uploaded_file.name.endswith(".docx"):
                    doc = docx.Document(save_path)
                    content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    st.write(f"✅ Extracted text from {len(doc.paragraphs)} paragraphs")
                else:
                    with open(save_path, "r", encoding="utf-8") as f:
                        content = f.read()
                    st.write(f"✅ Read {len(content)} characters")

                st.write("✂️ Splitting into chunks...")
                new_chunks = simple_chunk_text(content)
                st.write(f"✅ Created {len(new_chunks)} chunks")

                st.write("🧠 Generating embeddings and storing in ChromaDB...")
                existing_count = collection.count()
                new_ids = [f"chunk_{existing_count + i}" for i in range(len(new_chunks))]

                batch_size = 50
                progress_bar = st.progress(0)
                for i in range(0, len(new_chunks), batch_size):
                    batch_chunks = new_chunks[i:i+batch_size]
                    batch_ids = new_ids[i:i+batch_size]
                    collection.add(documents=batch_chunks, ids=batch_ids)
                    progress_bar.progress(min((i + batch_size) / len(new_chunks), 1.0))
                progress_bar.empty()
                st.write(f"✅ Indexed {len(new_chunks)} chunks in vector database")

                status.update(label=f"✅ '{uploaded_file.name}' added successfully!", state="complete", expanded=False)

    # ---------- Normal Chat Flow ----------
    if user_question:
        st.session_state.messages.append({"role": "user", "content": user_question})
        with st.chat_message("user"):
            st.write(user_question)

        with st.chat_message("assistant"):
            with st.spinner(f"Thinking using {model_choice}..."):
                answer = ask_policy_bot(user_question, model_choice)
                st.write(answer)
        st.session_state.messages.append({"role": "assistant", "content": answer})